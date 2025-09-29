#!/usr/bin/env python3
"""
Diagnostic script to test Word document import functionality.
This helps debug issues with heading detection and topic creation.

Usage: python debug_word_import.py path/to/document.docx
"""

import sys
import os
import re
from docx import Document

def detect_heading_level_from_style(style_name: str):
    """Extract heading level (1-6) from a Word style name."""
    if not style_name:
        return None
    sn = style_name.lower()
    patterns = [
        r'(?:^|[\s,;:()\-])heading\s+level\s*(\d)\b',   # heading level 2
        r'(?:^|[\s,;:()\-])sc\s+heading\s*(\d)\b',      # sc heading 2
        r'(?:^|[\s,;:()\-])heading\s*(\d)\b',            # heading 2
    ]
    for pat in patterns:
        m = re.search(pat, sn)
        if m:
            try:
                lvl = int(m.group(1))
                if 1 <= lvl <= 6:
                    return lvl
            except ValueError:
                return None
    return None

def analyze_word_document(file_path):
    """Analyze a Word document to show what headings would be detected."""
    print(f"🔍 Analyzing Word document: {file_path}")
    print("=" * 60)
    
    try:
        doc = Document(file_path)
        print(f"✅ Successfully opened document")
        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")
        print()
        
        headings_found = []
        topics_to_create = []
        current_content = []
        current_heading = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = (paragraph.text or '').strip()
            
            if not text:
                current_content.append("")  # Preserve empty lines for now
                continue
            
            # Get style information
            try:
                ps = getattr(paragraph, 'style', None)
                style_name_raw = getattr(ps, 'name', '') or ''
            except Exception:
                style_name_raw = ''
            
            style_name = style_name_raw.lower()
            
            # Check if this is a heading
            level = detect_heading_level_from_style(style_name)
            
            if level is not None:
                # This is a heading - commit previous topic if exists
                if current_heading and current_content:
                    content_text = '\n'.join(current_content).strip()
                    if content_text:  # Only if there's actual content
                        topics_to_create.append({
                            'title': current_heading,
                            'content': content_text,
                            'content_lines': len([line for line in current_content if line.strip()])
                        })
                
                # Start new topic
                current_heading = text
                current_content = []
                headings_found.append({
                    'paragraph_index': i,
                    'style_raw': style_name_raw,
                    'style_lower': style_name,
                    'detected_level': level,
                    'text': text
                })
                
                print(f"🎯 HEADING DETECTED:")
                print(f"   Paragraph {i}: '{text}'")
                print(f"   Style: '{style_name_raw}' -> Level {level}")
                print()
                
            else:
                # Regular content
                current_content.append(text)
                if i < 10:  # Show first 10 non-heading paragraphs for debugging
                    print(f"📝 Content paragraph {i}: '{text[:50]}...' (style: '{style_name_raw}')")
        
        # Don't forget the last topic
        if current_heading and current_content:
            content_text = '\n'.join(current_content).strip()
            if content_text:
                topics_to_create.append({
                    'title': current_heading,
                    'content': content_text,
                    'content_lines': len([line for line in current_content if line.strip()])
                })
        
        print()
        print("=" * 60)
        print("📊 ANALYSIS RESULTS:")
        print("=" * 60)
        
        print(f"🎯 Total headings found: {len(headings_found)}")
        print(f"📚 Topics that would be created: {len(topics_to_create)}")
        print()
        
        if headings_found:
            print("🎯 HEADINGS SUMMARY:")
            for h in headings_found:
                print(f"   • Level {h['detected_level']}: '{h['text']}' (style: {h['style_raw']})")
            print()
        
        if topics_to_create:
            print("📚 TOPICS SUMMARY:")
            for i, topic in enumerate(topics_to_create, 1):
                print(f"   {i}. '{topic['title']}'")
                print(f"      Content: {topic['content_lines']} lines, {len(topic['content'])} chars")
                if len(topic['content']) > 100:
                    preview = topic['content'][:100] + "..."
                else:
                    preview = topic['content']
                print(f"      Preview: {repr(preview)}")
                print()
        
        if not headings_found:
            print("❌ NO HEADINGS FOUND!")
            print("   This means no topics would be created during import.")
            print("   Common issues:")
            print("   • Document uses custom styles not matching 'Heading 1', 'Heading 2', etc.")
            print("   • Document uses formatting (bold, etc.) instead of heading styles")
            print("   • Document structure is not hierarchical")
            print()
            print("🔍 SHOWING ALL PARAGRAPH STYLES:")
            unique_styles = {}
            for p in doc.paragraphs:
                text = (p.text or '').strip()
                if text:
                    try:
                        style_name = getattr(p.style, 'name', 'Unknown') or 'Unknown'
                        if style_name not in unique_styles:
                            unique_styles[style_name] = []
                        unique_styles[style_name].append(text[:50])
                    except:
                        pass
            
            for style, examples in unique_styles.items():
                print(f"   • {style}: {len(examples)} paragraphs")
                if examples:
                    print(f"     Example: '{examples[0]}...'")
        
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")
        import traceback
        traceback.print_exc()

def main():
    if len(sys.argv) != 2:
        print("Usage: python debug_word_import.py path/to/document.docx")
        sys.exit(1)
    
    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        sys.exit(1)
    
    analyze_word_document(file_path)

if __name__ == "__main__":
    main()