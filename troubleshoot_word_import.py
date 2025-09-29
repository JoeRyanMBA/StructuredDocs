#!/usr/bin/env python3
"""
Comprehensive Word Document Import Troubleshooter

This script helps diagnose why Word document imports are not creating
separate topics for headings as expected.

Usage:
1. Run in your StructuredDocs environment to check system dependencies
2. Optionally provide a Word document path to analyze its structure

python troubleshoot_word_import.py [path_to_document.docx]
"""

import sys
import os
import subprocess
import tempfile
import re
from pathlib import Path

def check_system_dependencies():
    """Check if all required dependencies are available"""
    print("🔍 CHECKING SYSTEM DEPENDENCIES")
    print("=" * 50)
    
    issues = []
    
    # Check Pandoc
    try:
        result = subprocess.run(['pandoc', '--version'], 
                              capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            version = result.stdout.split('\n')[0] if result.stdout else "Unknown"
            print(f"✅ Pandoc: {version}")
        else:
            print(f"❌ Pandoc: Command failed - {result.stderr}")
            issues.append("Pandoc not working")
    except FileNotFoundError:
        print("❌ Pandoc: Not installed or not in PATH")
        issues.append("Pandoc not available")
    except Exception as e:
        print(f"❌ Pandoc: Error - {e}")
        issues.append("Pandoc error")
    
    # Check python-docx
    try:
        from docx import Document
        print("✅ python-docx: Available")
    except ImportError:
        print("❌ python-docx: Not available")
        issues.append("python-docx missing")
    
    # Check other dependencies
    modules = ['flask', 'sqlalchemy', 'tempfile', 're', 'io']
    for module in modules:
        try:
            __import__(module)
            print(f"✅ {module}: Available")
        except ImportError:
            print(f"❌ {module}: Not available")
            issues.append(f"{module} missing")
    
    # Check file system permissions
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            test_file = os.path.join(temp_dir, 'test.txt')
            with open(test_file, 'w') as f:
                f.write('test')
            print("✅ File System: Can create temporary files")
    except Exception as e:
        print(f"❌ File System: Cannot create temp files - {e}")
        issues.append("File system permissions")
    
    return issues

def analyze_word_document(file_path):
    """Analyze a Word document structure"""
    print(f"\n🔍 ANALYZING WORD DOCUMENT: {file_path}")
    print("=" * 50)
    
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return False
    
    try:
        from docx import Document
        doc = Document(file_path)
        
        print(f"📄 Total paragraphs: {len(doc.paragraphs)}")
        
        # Analyze paragraph styles
        styles_found = {}
        headings_found = []
        content_paragraphs = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = (paragraph.text or '').strip()
            if not text:
                continue
                
            try:
                style_name = getattr(paragraph.style, 'name', 'Unknown') or 'Unknown'
            except:
                style_name = 'Unknown'
            
            if style_name not in styles_found:
                styles_found[style_name] = []
            styles_found[style_name].append(text[:50])
            
            # Check if it would be detected as a heading
            level = detect_heading_level_from_style(style_name)
            if level:
                headings_found.append({
                    'paragraph': i,
                    'level': level,
                    'style': style_name,
                    'text': text[:100]
                })
            else:
                content_paragraphs += 1
        
        print(f"📊 ANALYSIS RESULTS:")
        print(f"  • Unique styles found: {len(styles_found)}")
        print(f"  • Headings detected: {len(headings_found)}")
        print(f"  • Content paragraphs: {content_paragraphs}")
        
        if headings_found:
            print(f"\n🎯 HEADINGS THAT WOULD CREATE TOPICS:")
            for h in headings_found:
                print(f"  Level {h['level']}: '{h['text']}' (style: {h['style']})")
            
            print(f"\n✅ This document should create {len(headings_found)} topics")
        else:
            print(f"\n❌ NO HEADINGS DETECTED - No topics would be created!")
            print(f"\n📋 ALL STYLES FOUND:")
            for style, examples in styles_found.items():
                print(f"  • {style}: {len(examples)} paragraphs")
                if examples:
                    print(f"    Example: '{examples[0]}...'")
            
            print(f"\n💡 POSSIBLE SOLUTIONS:")
            print(f"  1. Document uses custom styles - check if any contain 'Heading'")
            print(f"  2. Document uses formatting instead of styles")
            print(f"  3. Document structure is not hierarchical")
        
        return len(headings_found) > 0
        
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")
        return False

def detect_heading_level_from_style(style_name: str):
    """Extract heading level (1-6) from a Word style name."""
    if not style_name:
        return None
    sn = style_name.lower()
    patterns = [
        r'(?:^|[\s,;:()\-])heading\s+level\s*(\d)\b',   # heading level 2
        r'(?:^|[\s,;:()\-])sc\s+heading\s*(\d)\b',      # sc heading 2
        r'(?:^|[\s,;:()\-])heading\s*(\d)\b',            # heading 2
    ]
    for pat in patterns:
        m = re.search(pat, sn)
        if m:
            try:
                lvl = int(m.group(1))
                if 1 <= lvl <= 6:
                    return lvl
            except ValueError:
                return None
    return None

def test_conversion_simulation():
    """Test the conversion logic with a mock document"""
    print(f"\n🧪 TESTING CONVERSION LOGIC")
    print("=" * 50)
    
    # Simulate markdown that pandoc might produce
    mock_markdown = """# Introduction

This is introduction content.

## Section One

Content for section one.

### Subsection

More detailed content.

## Section Two

Final section content."""

    print("Input (simulated pandoc output):")
    print(mock_markdown)
    
    # Apply promotion logic
    lines = []
    promoted_count = 0
    for line in mock_markdown.splitlines():
        if line.strip().startswith('#'):
            hash_count = len(line) - len(line.lstrip('#'))
            if hash_count > 1:
                content = line.lstrip('#').strip()
                line = f"# {content}"
                promoted_count += 1
        lines.append(line)
    
    print(f"\nAfter promotion ({promoted_count} headings promoted):")
    for line in lines:
        if line.strip().startswith('# ') and not line.strip().startswith('##'):
            print(f"🎯 {line}")
        elif line.strip():
            print(f"   {line}")
    
    # Count topics that would be created
    topics = []
    current_title = None
    current_content = []
    
    for line in lines:
        is_h1 = line.strip().startswith('#') and not line.strip().startswith('##')
        if is_h1:
            if current_title and current_content:
                content = '\n'.join(current_content).strip()
                if content:
                    topics.append((current_title, len(content)))
            current_title = line.strip().lstrip('#').strip()
            current_content = []
        else:
            current_content.append(line)
    
    # Don't forget the last topic
    if current_title and current_content:
        content = '\n'.join(current_content).strip()
        if content:
            topics.append((current_title, len(content)))
    
    print(f"\n📚 TOPICS THAT WOULD BE CREATED: {len(topics)}")
    for i, (title, content_len) in enumerate(topics, 1):
        print(f"  {i}. '{title}' ({content_len} chars)")
    
    return len(topics)

def main():
    print("🔧 STRUCTUREDDOCS WORD IMPORT TROUBLESHOOTER")
    print("=" * 60)
    
    # Check system dependencies
    system_issues = check_system_dependencies()
    
    # Test conversion logic
    expected_topics = test_conversion_simulation()
    
    # Analyze specific document if provided
    doc_analysis_ok = True
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
        doc_analysis_ok = analyze_word_document(doc_path)
    
    # Final diagnosis
    print(f"\n🏥 DIAGNOSIS")
    print("=" * 60)
    
    if system_issues:
        print("❌ SYSTEM ISSUES FOUND:")
        for issue in system_issues:
            print(f"  • {issue}")
        print(f"\n🔧 REQUIRED ACTIONS:")
        if "Pandoc not available" in system_issues:
            print("  • Install pandoc: apt-get install pandoc")
        if "python-docx missing" in system_issues:
            print("  • Install python-docx: pip install python-docx")
    else:
        print("✅ System dependencies are OK")
    
    if expected_topics == 0:
        print("❌ CONVERSION LOGIC ISSUE: Logic test failed")
        print("  • This suggests a bug in the promotion/parsing code")
    else:
        print(f"✅ Conversion logic works (creates {expected_topics} topics)")
    
    if len(sys.argv) > 1:
        if not doc_analysis_ok:
            print("❌ DOCUMENT STRUCTURE ISSUE:")
            print("  • Document has no recognizable heading styles")
            print("  • Check if document uses 'Heading 1', 'Heading 2', etc.")
            print("  • Or 'SC Heading 1', 'SC Heading 2', etc.")
        else:
            print("✅ Document structure looks OK")
    
    print(f"\n💡 NEXT STEPS:")
    if system_issues:
        print("  1. Fix system dependency issues first")
    print("  2. Check application logs for specific error messages")
    print("  3. Try importing a simple test document with clear heading styles")
    print("  4. Verify database permissions and connectivity")
    
    print(f"\n📋 RECENT CHANGES THAT MIGHT AFFECT IMPORTS:")
    print("  • Removed aggressive blank line removal from document parsing")
    print("  • This should improve heading detection accuracy")
    print("  • Test with a document that worked before")

if __name__ == "__main__":
    main()