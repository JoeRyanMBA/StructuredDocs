from io import BytesIO
from types import SimpleNamespace

from docx import Document
from flask import Flask

from backend.routes.import_handler import (
    _parse_and_store,
    _parse_hierarchical_content,
    _convert_docx_to_markdown_fallback,
    _convert_word_to_markdown_no_images,
)
from backend.models import ImportItem


def _build_docx_bytes():
    doc = Document()
    doc.add_heading('Main Topic', level=1)
    doc.add_paragraph('Intro paragraph')
    doc.add_paragraph('List bullet', style='List Bullet')
    doc.add_paragraph('Nested bullet', style='List Bullet 2')
    doc.add_paragraph('List bullet 2', style='List Bullet')
    doc.add_paragraph('List number', style='List Number')
    doc.add_paragraph('Nested number', style='List Number 2')
    doc.add_paragraph('List number 2', style='List Number')

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_docx_fallback_preserves_word_list_markers():
    markdown = _convert_docx_to_markdown_fallback(_build_docx_bytes())

    assert '# Main Topic' in markdown
    assert 'Intro paragraph' in markdown
    assert '- List bullet' in markdown
    assert '    - Nested bullet' in markdown
    assert '- List bullet 2' in markdown
    assert '1. List number' in markdown
    assert '    1. Nested number' in markdown
    assert '2. List number 2' in markdown


def test_docx_pandoc_conversion_targets_gfm(monkeypatch):
    captured = {}

    def fake_run(cmd, capture_output, text, timeout):
        captured['cmd'] = cmd
        output_path = cmd[cmd.index('-o') + 1]
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('| Col A | Col B |\n| --- | --- |\n| 1 | 2 |\n')
        return SimpleNamespace(returncode=0, stderr='')

    monkeypatch.setattr('backend.routes.import_handler.subprocess.run', fake_run)

    app = Flask(__name__)
    with app.app_context():
        markdown = _convert_word_to_markdown_no_images(b'fake-docx-payload')

    assert '--to' in captured['cmd']
    assert captured['cmd'][captured['cmd'].index('--to') + 1] == 'gfm'
    assert '| Col A | Col B |' in markdown


def test_parse_and_store_preserves_docx_table_markdown(monkeypatch):
    markdown_with_table = """# Safety Checklist
Intro paragraph before table.

| Item | Owner |
| --- | --- |
| Lockout | Maintenance |
"""

    monkeypatch.setattr(
        'backend.routes.import_handler._convert_word_to_markdown',
        lambda file_content, import_doc_id: markdown_with_table,
    )
    monkeypatch.setattr('backend.routes.import_handler._extract_and_store_links', lambda document_id, content: 0)

    captured_items = []

    def fake_add(obj):
        if isinstance(obj, ImportItem):
            captured_items.append(obj)

    monkeypatch.setattr('backend.routes.import_handler.db.session.add', fake_add)

    app = Flask(__name__)
    with app.app_context():
        class DummyFile:
            def __init__(self, payload):
                self.stream = BytesIO(payload)

            def read(self):
                return self.stream.read()

        dummy_file = DummyFile(b'fake-docx-binary')
        dummy_imp_doc = SimpleNamespace(id=42, filename='safety.docx')
        _parse_and_store(dummy_file, dummy_imp_doc, source='word', preserve_hierarchy=False)

    assert len(captured_items) == 1
    item = captured_items[0]
    assert item.title == 'Safety Checklist'
    assert '| Item | Owner |' in item.content
    assert '| --- | --- |' in item.content
    assert '| Lockout | Maintenance |' in item.content


def test_parse_hierarchical_content_ignores_blank_heading_lines():
    markdown_with_blank_heading = """# Safety Checklist
Intro paragraph.

#

More content after the break.
"""

    app = Flask(__name__)
    with app.app_context():
        items = _parse_hierarchical_content(markdown_with_blank_heading)

    assert len(items) == 1
    assert items[0]['title'] == 'Safety Checklist'
    assert 'Intro paragraph.' in items[0]['content']
    assert 'More content after the break.' in items[0]['content']
