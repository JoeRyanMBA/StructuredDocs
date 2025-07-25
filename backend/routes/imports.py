# backend/routes/imports.py

from flask import Blueprint, request, jsonify, current_app, make_response
from werkzeug.utils import secure_filename
from backend.models import db, ImportDocument, ImportItem, Topic
from docx import Document as DocxDocument
import re

imports = Blueprint('imports', __name__, url_prefix='/api/import')
SOURCES = ('word', 'markdown')


def _parse_and_store(file, imp_doc, source):
    """Extract heading‐1s and content blocks into ImportItem rows."""
    file.stream.seek(0)

    if source == 'word':
        doc = DocxDocument(file)
        paras = [(p.style.name, p.text) for p in doc.paragraphs]
    else:
        raw = file.read().decode('utf-8')
        paras = [('md', line) for line in raw.splitlines()]

    items, buffer, order, current_title = [], [], 0, None

    def commit_buffer():
        nonlocal order, current_title, buffer
        if current_title:
            content = '\n'.join(buffer).strip()
            items.append((order, current_title, content))
            order += 1
            buffer = []

    for style, text in paras:
        is_h1 = (
            source == 'word' and style.startswith('Heading 1')
        ) or (
            source == 'markdown' and re.match(r'^\s*#\s+', text)
        )
        if is_h1:
            commit_buffer()
            current_title = (
                text.strip() if source == 'word'
                else re.sub(r'^\s*#\s+', '', text).strip()
            )
        else:
            buffer.append(text)

    commit_buffer()

    for order, title, content in items:
        db.session.add(ImportItem(
            document_id=imp_doc.id,
            heading_order=order,
            title=title,
            content=content
        ))


def _upload_file(source):
    file = request.files.get('file')
    if not file or source not in SOURCES:
        return jsonify({'error': 'Missing file or invalid source'}), 400

    try:
        imp_doc = ImportDocument(
            filename=secure_filename(file.filename),
            source_type=source
        )
        db.session.add(imp_doc)
        db.session.flush()  # get imp_doc.id

        _parse_and_store(file, imp_doc, source)
        db.session.commit()
        return jsonify(imp_doc.to_dict(include_items=True)), 201

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("Upload failed")
        return jsonify({'error': str(e)}), 500


@imports.route('/upload', methods=['POST'])
def upload_generic():
    return _upload_file(request.form.get('source', '').lower())


@imports.route('/markdown', methods=['POST'])
def upload_markdown():
    return _upload_file('markdown')


@imports.route('/word', methods=['POST'])
def upload_word():
    return _upload_file('word')


@imports.route('/staging/<int:doc_id>', methods=['GET'])
def get_staging(doc_id):
    doc = ImportDocument.query.get_or_404(doc_id)
    return jsonify(doc.to_dict(include_items=True)), 200


@imports.route('/staging/<int:doc_id>/sme_approve', methods=['POST'])
def sme_approve(doc_id):
    doc = ImportDocument.query.get_or_404(doc_id)
    if doc.review_step != 'pending':
        return jsonify({'error': 'Already reviewed'}), 400

    doc.review_step = 'sme_approved'
    doc.reviewer    = request.headers.get('X-User', 'SME')
    doc.reviewed_at = db.func.now()
    db.session.commit()
    return jsonify(doc.to_dict()), 200


@imports.route('/staging/<int:doc_id>/commit', methods=['POST'])
def commit_import(doc_id):
    doc = ImportDocument.query.get_or_404(doc_id)
    if doc.review_step != 'sme_approved':
        return jsonify({'error': 'Must SME-approve first'}), 400

    try:
        for item in doc.items:
            topic = Topic(
                title=item.title,
                content=item.content,
                status='draft'
            )
            db.session.add(topic)
            db.session.flush()
            item.committed_topic = topic.id

        doc.status      = 'approved'
        doc.review_step = 'final_approved'
        doc.reviewed_at = db.func.now()
        db.session.commit()
        return jsonify({'message': 'Imported', 'import_id': doc.id}), 200

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("Commit failed")
        return jsonify({'error': str(e)}), 500


@imports.route('/staging/<int:doc_id>/reject', methods=['POST'])
def reject_import(doc_id):
    doc = ImportDocument.query.get_or_404(doc_id)
    doc.status      = 'rejected'
    doc.reviewed_at = db.func.now()
    db.session.commit()
    return jsonify({'message': 'Rejected'}), 200


@imports.route('/staging/<int:doc_id>/export', methods=['GET'])
def export_import(doc_id):
    doc = ImportDocument.query.get_or_404(doc_id)

    lines = []
    for item in doc.items:
        lines.append(f"# {item.title}")
        lines.append(item.content or '')
        lines.append('')
    md = "\n".join(lines)

    resp = make_response(md)
    resp.headers['Content-Type'] = 'text/markdown'

    fn = doc.filename
    if not fn.lower().endswith('.md'):
        fn += '.md'
    resp.headers['Content-Disposition'] = f'attachment; filename="{fn}"'
    return resp


@imports.route('/history', methods=['GET'])
def import_history():
    docs = ImportDocument.query.order_by(
        ImportDocument.created_at.desc()
    ).all()
    return jsonify([d.to_dict() for d in docs]), 200