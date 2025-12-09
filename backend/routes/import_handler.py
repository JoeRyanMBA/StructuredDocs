from flask import Blueprint, request, jsonify, current_app, make_response
from werkzeug.utils import secure_filename
from ..models import db, ImportDocument, ImportItem, ImportImage, ImportLink, Topic
from ..utils.image_handler import ImageHandler
import re
from docx import Document
import io
import subprocess
import tempfile
import os
import shutil
import uuid
from urllib.parse import urlparse

import_bp = Blueprint('import_handler', __name__, url_prefix='/api/import')
SOURCES = ('word', 'markdown')

def detect_heading_level_from_style(style_name: str):
    """Extract heading level (1-6) from a Word style name.

    Supports variants like:
      - Heading 1
      - SC Heading 2 / Heading 2, SC Heading 2
      - Heading 3 (SC)
      - Heading Level 4
      - HEADING LEVEL 5 (case-insensitive)
    Returns int or None if not matched.
    """
    if not style_name:
        return None
    sn = style_name.lower()
    patterns = [
        r'(?:^|[\s,;:()\-])heading\s+level\s*(\d)\b',   # heading level 2
        r'(?:^|[\s,;:()\-])sc\s+heading\s*(\d)\b',      # sc heading 2
        r'(?:^|[\s,;:()\-])heading\s*(\d)\b',            # heading 2
    ]
    for pat in patterns:
        m = re.search(pat, sn)
        if m:
            try:
                lvl = int(m.group(1))
                if 1 <= lvl <= 6:
                    return lvl
            except ValueError:
                return None
    return None


def _convert_word_to_markdown(file_content, import_doc_id):
    """Convert Word document to Markdown using pandoc with proper image handling"""
    try:
        # Create unique temporary directory for this import
        temp_base_dir = tempfile.mkdtemp(prefix=f'import_{import_doc_id}_')
        temp_media_dir = os.path.join(temp_base_dir, 'media')
        
        # Create temporary files for input and output
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, dir=temp_base_dir) as temp_input:
            temp_input.write(file_content)
            temp_input_path = temp_input.name
        
        with tempfile.NamedTemporaryFile(suffix='.md', delete=False, dir=temp_base_dir) as temp_output:
            temp_output_path = temp_output.name
        
        try:
            # Use pandoc to convert Word to Markdown with better list handling
            cmd = [
                'pandoc',
                '--from', 'docx',
                '--to', 'markdown',
                '--wrap', 'none',  # Don't wrap lines
                '--extract-media', temp_media_dir,  # Extract images to our temp media dir
                '--markdown-headings=atx',  # Use ATX-style headings (#)
                '--list-tables',  # Use pipe tables for better compatibility
                '--strip-comments',  # Remove HTML comments
                temp_input_path,
                '-o', temp_output_path
            ]
            
            print(f"PANDOC: Running command: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                print(f"PANDOC ERROR: {result.stderr}")
                raise Exception(f"Pandoc conversion failed: {result.stderr}")
            
            # Read the converted Markdown
            with open(temp_output_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            print(f"PANDOC SUCCESS: Converted {len(file_content)} bytes to {len(markdown_content)} chars of Markdown")
            
            # Initialize image handler
            image_handler = ImageHandler(import_doc_id)
            
            # Extract and store images permanently
            updated_markdown, stored_images = image_handler.extract_and_store_images(
                temp_media_dir, markdown_content
            )
            
            # Store image metadata in database
            for image_info in stored_images:
                import_image = ImportImage(
                    document_id=import_doc_id,
                    filename=image_info['filename'],
                    original_name=image_info['original_name'],
                    public_url=image_info['public_url'],
                    backend_path=image_info['backend_path'],
                    frontend_path=image_info['frontend_path'],
                    width=image_info['width'],
                    height=image_info['height'],
                    format=image_info['format'],
                    file_size=image_info['file_size'],
                    mime_type=image_info['mime_type']
                )
                db.session.add(import_image)
            
            print(f"IMAGE PROCESSING: Stored {len(stored_images)} images")
            
            # Post-process the markdown to fix issues
            updated_markdown = _post_process_markdown(updated_markdown)
            
            # Additional cleaning for HTML comments and formatting issues
            updated_markdown = _clean_markdown_content(updated_markdown)
            
            # Fix progressive list indentation issues from Word conversion
            updated_markdown = _fix_list_indentation(updated_markdown)
            print(f"LIST INDENTATION: Fixed progressive indentation issues")
            
            # NOTE: Blank line removal now happens per-topic during content cleaning
            # to preserve document structure for heading detection
            
            # Validate image references
            validation_issues = image_handler.validate_markdown_images(updated_markdown)
            if validation_issues:
                for issue in validation_issues:
                    print(f"IMAGE VALIDATION: {issue['message']}")
            
            # Clean up temporary directory
            image_handler.cleanup_temp_images(temp_base_dir)
            
            return updated_markdown
            
        finally:
            # Clean up temporary files
            try:
                os.unlink(temp_input_path)
                os.unlink(temp_output_path)
            except OSError:
                pass  # Files might already be deleted
                
    except subprocess.TimeoutExpired:
        print("PANDOC ERROR: Conversion timed out")
        raise Exception("Word to Markdown conversion timed out")
    except FileNotFoundError:
        print("PANDOC ERROR: Pandoc not found")
        raise Exception("Pandoc is not installed. Please contact your system administrator to install pandoc for Word document conversion.")
    except subprocess.CalledProcessError as e:
        print(f"PANDOC ERROR: Process failed with code {e.returncode}: {e.stderr}")
        raise Exception(f"Pandoc conversion failed: {e.stderr}")
    except Exception as e:
        print(f"PANDOC ERROR: {str(e)}")
        # Check if it's a pandoc not found error
        if "pandoc" in str(e).lower() and ("not found" in str(e).lower() or "command not found" in str(e).lower()):
            raise Exception("Pandoc is not installed. Please contact your system administrator to install pandoc for Word document conversion.")
        raise Exception(f"Failed to convert Word to Markdown: {str(e)}")


def _convert_word_to_markdown_no_images(file_content):
    """Convert Word document to Markdown using pandoc without image processing"""
    try:
        # Create unique temporary directory for this conversion
        temp_base_dir = tempfile.mkdtemp(prefix='temp_conversion_')
        
        # Create temporary files for input and output
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, dir=temp_base_dir) as temp_input:
            temp_input.write(file_content)
            temp_input_path = temp_input.name
        
        with tempfile.NamedTemporaryFile(suffix='.md', delete=False, dir=temp_base_dir) as temp_output:
            temp_output_path = temp_output.name
        
        try:
            # Use pandoc to convert Word to Markdown without image extraction
            cmd = [
                'pandoc',
                '--from', 'docx',
                '--to', 'markdown',
                '--wrap', 'none',
                '--markdown-headings=atx',
                '--list-tables',
                '--strip-comments',
                temp_input_path,
                '-o', temp_output_path
            ]
            
            print(f"PANDOC (NO IMAGES): Running command: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
            
            if result.returncode != 0:
                print(f"PANDOC ERROR: {result.stderr}")
                raise Exception(f"Pandoc conversion failed: {result.stderr}")
            
            # Read the converted Markdown
            with open(temp_output_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            print(f"PANDOC SUCCESS (NO IMAGES): Converted {len(file_content)} bytes to {len(markdown_content)} chars")
            
            return markdown_content
            
        finally:
            # Clean up temporary files
            try:
                os.unlink(temp_input_path)
                os.unlink(temp_output_path)
                shutil.rmtree(temp_base_dir)
            except Exception as cleanup_error:
                print(f"Cleanup warning: {cleanup_error}")
                
    except subprocess.TimeoutExpired:
        print("PANDOC ERROR: Conversion timed out")
        raise Exception("Word document conversion timed out")
    except FileNotFoundError:
        print("PANDOC ERROR: Pandoc not found")
        raise Exception("Pandoc is not installed. Please contact your system administrator to install pandoc for Word document conversion.")
    except subprocess.CalledProcessError as e:
        print(f"PANDOC ERROR: Process failed with code {e.returncode}: {e.stderr}")
        raise Exception(f"Pandoc conversion failed: {e.stderr}")
    except Exception as e:
        print(f"PANDOC ERROR: {str(e)}")
        # Check if it's a pandoc not found error
        if "pandoc" in str(e).lower() and ("not found" in str(e).lower() or "command not found" in str(e).lower()):
            raise Exception("Pandoc is not installed. Please contact your system administrator to install pandoc for Word document conversion.")
        raise Exception(f"Failed to convert Word to Markdown: {str(e)}")


def _post_process_markdown(markdown_content):
    """Post-process markdown to fix nested lists and handle margin notes while preserving blank lines"""
    lines = markdown_content.split('\n')
    processed_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle HTML comment blocks that contain nested lists
        if line.strip() == '```{=html}' and i + 1 < len(lines):
            # Look for the closing block
            j = i + 1
            html_content = []
            while j < len(lines) and lines[j].strip() != '```':
                html_content.append(lines[j])
                j += 1
            
            if j < len(lines):  # Found closing ```
                # Process the HTML content to extract nested lists
                html_text = '\n'.join(html_content)
                converted_list = _convert_html_list_to_markdown(html_text)
                if converted_list:
                    # Replace the entire HTML fenced block with converted markdown
                    processed_lines.extend(converted_list.split('\n'))
                else:
                    # Could not convert; keep original fenced block
                    processed_lines.append(lines[i])
                    processed_lines.extend(html_content)
                    processed_lines.append('```')
                # Skip past the closing fence
                i = j + 1
                continue
            else:
                # No closing fence found; keep the line as-is
                processed_lines.append(line)
                i += 1
                continue
        
        # Default: keep line
        processed_lines.append(line)
        i += 1

    # Join back the processed lines; further cleaning handled by other steps
    return '\n'.join(processed_lines)


def _convert_html_list_to_markdown(html_content):
    """Convert HTML list content to proper markdown lists"""
    # Remove HTML comments
    html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
    
    # Convert <ul> and <li> tags to markdown
    # This is a simple conversion - for complex nested lists you might need a proper HTML parser
    html_content = re.sub(r'<ul[^>]*>', '', html_content)
    html_content = re.sub(r'</ul>', '', html_content)
    html_content = re.sub(r'<ol[^>]*>', '', html_content)
    html_content = re.sub(r'</ol>', '', html_content)
    
    lines = html_content.split('\n')
    markdown_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Convert <li> tags to markdown list items
        if line.startswith('<li>') and line.endswith('</li>'):
            content = line[4:-5].strip()  # Remove <li> and </li>
            # Determine nesting level based on leading whitespace in original
            # For now, just use simple bullet points
            markdown_lines.append(f"- {content}")
        elif '<li>' in line:
            # Handle multi-line or partial li tags
            content = re.sub(r'</?li[^>]*>', '', line).strip()
            if content:
                markdown_lines.append(f"- {content}")
    
    return '\n'.join(markdown_lines) if markdown_lines else None


def _clean_markdown_content(content):
    """Additional cleaning for problematic patterns in converted markdown"""
    # Remove HTML comments that appear as standalone lines or inline
    content = re.sub(r'<!--\s*-->\s*\n?', '', content)
    content = re.sub(r'<!---->\s*\n?', '', content)
    
    # Fix patterns where list items are separated by HTML comments
    # Pattern: list item, HTML comment, list item -> continuous list
    lines = content.split('\n')
    cleaned_lines = []
    
    for i, line in enumerate(lines):
        # Skip HTML comment lines entirely
        if re.match(r'^\s*<!--.*?-->\s*$', line):
            continue
        
        # Skip empty paragraphs (lines with only whitespace)
        if line.strip() == '':
            # Keep one blank line but skip excessive ones
            if cleaned_lines and cleaned_lines[-1].strip() != '':
                cleaned_lines.append('')
            continue
        
        # Skip lines that are just paragraph tags or whitespace
        if re.match(r'^\s*</?p>\s*$', line):
            continue
            
        # Clean up lines that contain only HTML paragraph tags with whitespace
        line = re.sub(r'^\s*<p>\s*</p>\s*$', '', line)
        if line.strip() == '':
            continue
        
        # Remove empty markdown paragraph indicators
        if re.match(r'^\s*&nbsp;\s*$', line):
            continue
            
        # Remove lines with only whitespace characters and HTML entities
        if re.match(r'^\s*(&nbsp;|\s|&\w+;)*\s*$', line):
            continue
        
        # Keep the line
        cleaned_lines.append(line)
    
    # Join back and clean up excessive whitespace
    cleaned_content = '\n'.join(cleaned_lines)
    
    # Remove excessive blank lines (more than 2 consecutive)
    cleaned_content = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned_content)
    
    # Remove blank lines at the beginning and end
    cleaned_content = cleaned_content.strip()
    
    # Remove empty list items that might have been created
    cleaned_content = re.sub(r'^[\s]*[-*+]\s*$', '', cleaned_content, flags=re.MULTILINE)
    
    # Clean up any remaining empty paragraph patterns
    cleaned_content = re.sub(r'\n\s*<p>\s*</p>\s*\n', '\n', cleaned_content)
    cleaned_content = re.sub(r'<p>\s*</p>', '', cleaned_content)
    
    return cleaned_content


def _extract_and_store_links(document_id, content, position=0):
    """Extract links from content and store in database
    
    Args:
        document_id: The ImportDocument ID
        content: The text content to extract links from
        position: Starting position in document for link ordering
        
    Returns:
        int: Number of links extracted
    """
    if not content:
        return 0
    
    links_extracted = 0
    
    # Pattern to match markdown links: [text](url) or [text](url "title")
    markdown_link_pattern = r'\[([^\]]+)\]\(([^)]+?)(?:\s+"([^"]*)")?\)'
    
    # Pattern to match plain URLs (http/https)
    url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+[^\s<>"{}|\\^`\[\].,;:!?)]'
    
    # Extract markdown-style links
    for match in re.finditer(markdown_link_pattern, content):
        title = match.group(1).strip()
        url = match.group(2).strip()
        description = match.group(3).strip() if match.group(3) else None
        
        # Truncate title to 200 characters (database column limit)
        if len(title) > 200:
            title = title[:197] + '...'
        
        # Get surrounding context (50 chars before and after)
        start = max(0, match.start() - 50)
        end = min(len(content), match.end() + 50)
        context = content[start:end].replace('\n', ' ').strip()
        
        # Determine link type and if it's internal
        link_type, is_internal = _classify_link(url)
        
        # Store the link
        import_link = ImportLink(
            document_id=document_id,
            title=title,
            url=url,
            description=description,
            link_type=link_type,
            is_internal=is_internal,
            context=context,
            position_in_document=position + match.start()
        )
        
        db.session.add(import_link)
        links_extracted += 1
        print(f"LINK EXTRACTED: [{title}]({url}) - type: {link_type}, internal: {is_internal}")
    
    # Extract plain URLs (not already captured as markdown links)
    existing_urls = set()
    for match in re.finditer(markdown_link_pattern, content):
        existing_urls.add(match.group(2).strip())
    
    for match in re.finditer(url_pattern, content):
        url = match.group(0)
        if url in existing_urls:
            continue  # Skip URLs already captured as markdown links
        
        # Create title from URL
        parsed = urlparse(url)
        title = parsed.netloc or url
        
        # Truncate title to 200 characters (database column limit)
        if len(title) > 200:
            title = title[:197] + '...'
        
        # Get surrounding context
        start = max(0, match.start() - 50)
        end = min(len(content), match.end() + 50)
        context = content[start:end].replace('\n', ' ').strip()
        
        # Determine link type and if it's internal
        link_type, is_internal = _classify_link(url)
        
        # Store the link
        import_link = ImportLink(
            document_id=document_id,
            title=title,
            url=url,
            description=None,
            link_type=link_type,
            is_internal=is_internal,
            context=context,
            position_in_document=position + match.start()
        )
        
        db.session.add(import_link)
        links_extracted += 1
        print(f"URL EXTRACTED: {url} - type: {link_type}, internal: {is_internal}")
    
    return links_extracted


def _classify_link(url):
    """Classify a link by type and determine if it's internal
    
    Returns:
        tuple: (link_type, is_internal)
    """
    if not url:
        return 'other', False
    
    url_lower = url.lower()
    parsed = urlparse(url)
    
    # Determine if internal (same domain or relative URL)
    is_internal = False
    if not parsed.netloc:  # Relative URL
        is_internal = True
    elif parsed.netloc:
        # You could add logic here to check if it's your internal domain
        # For now, assume external unless it's a relative path
        is_internal = False
    
    # Classify link type based on URL patterns and file extensions
    if any(word in url_lower for word in ['form', 'submit', 'application']):
        return 'form', is_internal
    elif any(ext in url_lower for ext in ['.pdf', '.doc', '.docx', '.txt', '.rtf']):
        return 'document', is_internal
    elif any(word in url_lower for word in ['policy', 'policies']):
        return 'policy', is_internal
    elif any(word in url_lower for word in ['procedure', 'process', 'sop']):
        return 'procedure', is_internal
    elif any(word in url_lower for word in ['regulation', 'rule', 'law', 'legal']):
        return 'regulation', is_internal
    elif parsed.netloc:  # Has domain, likely external website
        return 'website', is_internal
    else:
        return 'other', is_internal


def _fix_list_indentation(content):
    """Fix progressive list indentation issues from Word document conversion"""
    if not content or not content.strip():
        return content
    
    lines = content.split('\n')
    fixed_lines = []
    
    # Track the current indentation context
    current_list_type = None  # 'bullet' or 'numbered'
    indentation_stack = []  # Track indentation levels
    
    for line in lines:
        original_line = line
        stripped = line.strip()
        
        # Check if this is a list item (bullet or numbered)
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.*)$', line)
        numbered_match = re.match(r'^(\s*)(\d+\.)\s+(.*)$', line)
        
        if bullet_match:
            leading_spaces = bullet_match.group(1)
            content_text = bullet_match.group(2)
            space_count = len(leading_spaces)
            
            # Determine indentation level (0, 1, 2, or 3 max)
            if space_count == 0:
                level = 0
            elif space_count <= 3:
                level = 1
            elif space_count <= 7:
                level = 2
            else:
                level = 3  # Cap at 3 levels
            
            # Apply consistent indentation
            indent = "  " * level  # 2 spaces per level
            fixed_line = f"{indent}- {content_text}"
            fixed_lines.append(fixed_line)
            
        elif numbered_match:
            leading_spaces = numbered_match.group(1)
            number_part = numbered_match.group(2)
            content_text = numbered_match.group(3)
            space_count = len(leading_spaces)
            
            # Determine indentation level (0, 1, 2, or 3 max)
            if space_count == 0:
                level = 0
            elif space_count <= 3:
                level = 1
            elif space_count <= 7:
                level = 2
            else:
                level = 3  # Cap at 3 levels
            
            # Apply consistent indentation
            indent = "  " * level  # 2 spaces per level
            fixed_line = f"{indent}{number_part} {content_text}"
            fixed_lines.append(fixed_line)
            
        else:
            # Not a list item, keep as is
            fixed_lines.append(original_line)
    
    return '\n'.join(fixed_lines)


def _remove_all_blank_lines(content):
    """Remove all blank lines from content - specifically for Word document imports"""
    if not content or not content.strip():
        return ''
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Skip any line that is empty or contains only whitespace
        if line.strip() == '':
            continue
        
        # Skip lines that are just HTML paragraph tags (including with whitespace inside)
        if re.match(r'^\s*<\s*/?p\s*/?>\s*$', line):
            continue
        
        # Skip empty HTML paragraph pairs specifically
        if re.match(r'^\s*<p></p>\s*$', line):
            continue
        
        # Skip lines with only non-breaking spaces or similar HTML entities
        if re.match(r'^\s*(&nbsp;|\s|&\w+;)*\s*$', line):
            continue
            
        # Skip empty list items
        if re.match(r'^\s*[-*+]\s*$', line):
            continue
            
        # Skip empty numbered list items
        if re.match(r'^\s*\d+\.\s*$', line):
            continue
        
        # Skip Word document artifacts
        if re.match(r'^\s*\[\s*\]\s*$', line):  # Empty checkboxes
            continue
            
        if re.match(r'^\s*\\\s*$', line):  # Stray backslashes
            continue
            
        # Skip lines with only formatting marks or tabs
        if re.match(r'^\s*[\t\r\f\v]+\s*$', line):
            continue
            
        # Skip Word table artifacts like empty table cells
        if re.match(r'^\s*\|\s*\|\s*$', line):
            continue
        
        # Skip HTML comments
        if re.match(r'^\s*<!--.*?-->\s*$', line):
            continue
        
        # Skip empty HTML tags more broadly
        if re.match(r'^\s*<\s*/?[^>]*>\s*$', line.strip()):
            continue
        
        # Keep all non-empty lines
        cleaned_lines.append(line)
    
    # Join lines without any blank lines between them
    cleaned_content = '\n'.join(cleaned_lines)
    
    # Remove leading/trailing whitespace
    cleaned_content = cleaned_content.strip()
    
    return cleaned_content


def _clean_topic_content(content):
    """Clean up topic content by removing empty paragraphs and excessive whitespace"""
    if not content or not content.strip():
        return ''
    
    lines = content.split('\n')
    cleaned_lines = []
    
    for line in lines:
        # Skip completely empty lines, but preserve intentional spacing
        if line.strip() == '':
            # Only add blank line if the previous line wasn't blank
            if cleaned_lines and cleaned_lines[-1].strip() != '':
                cleaned_lines.append('')
            continue
        
        # Remove lines that are just HTML paragraph tags
        if re.match(r'^\s*</?p>\s*$', line):
            continue
        
        # Remove lines with only non-breaking spaces or similar
        if re.match(r'^\s*(&nbsp;|\s|&\w+;)*\s*$', line):
            continue
            
        # Remove empty list items
        if re.match(r'^\s*[-*+]\s*$', line):
            continue
            
        # Remove empty numbered list items
        if re.match(r'^\s*\d+\.\s*$', line):
            continue
        
        # Remove Word document artifacts
        if re.match(r'^\s*\[\s*\]\s*$', line):  # Empty checkboxes
            continue
            
        if re.match(r'^\s*\\\s*$', line):  # Stray backslashes
            continue
            
        # Remove lines with only formatting marks or tabs
        if re.match(r'^\s*[\t\r\f\v]+\s*$', line):
            continue
            
        # Remove Word table artifacts like empty table cells
        if re.match(r'^\s*\|\s*\|\s*$', line):
            continue
        
        cleaned_lines.append(line)
    
    # Join and clean up
    cleaned_content = '\n'.join(cleaned_lines)
    
    # Remove excessive blank lines (more than 2 consecutive)
    cleaned_content = re.sub(r'\n\s*\n\s*\n\s*\n+', '\n\n\n', cleaned_content)
    
    # Clean up common Word artifacts in the content
    cleaned_content = re.sub(r'\s*\n\s*\n\s*\n\s*\n+', '\n\n', cleaned_content)  # Multiple newlines
    # NOTE: Removed the aggressive single newline joining to preserve paragraph breaks
    
    # Remove leading/trailing whitespace
    cleaned_content = cleaned_content.strip()
    
    # If content is now just whitespace or empty, return empty string
    if not cleaned_content or not cleaned_content.strip():
        return ''
    
    return cleaned_content


def _parse_and_store(file, imp_doc, source, preserve_hierarchy=False):
    """Extract heading‐1s and content blocks into ImportItem rows."""
    file.stream.seek(0)

    if source == 'word':
        # Parse Word document: try pandoc first, then python-docx fallback
        print(f"PARSING WORD DOC: {imp_doc.filename}")
        # Read file content for conversion attempts
        file_content = file.read()
        file.stream.seek(0)  # Reset stream for potential future reads

        lines = []
        heading_levels = []  # Track original heading levels for hierarchy
        try:
            # Primary conversion via pandoc
            markdown_content = _convert_word_to_markdown(file_content, imp_doc.id)

            # Parse lines and optionally promote headers based on preserve_hierarchy setting
            for line in markdown_content.splitlines():
                if line.strip().startswith('#'):
                    hash_count = len(line) - len(line.lstrip('#'))
                    heading_levels.append(hash_count)  # Store original level
                    
                    if preserve_hierarchy:
                        # Keep original heading level
                        lines.append(line)
                        print(f"HEADING PRESERVED: '{line.strip()}' (level H{hash_count})")
                    else:
                        # Promote all headings to H1 (original behavior)
                        heading_text = line.lstrip('#').strip()
                        promoted_line = f"# {heading_text}"
                        lines.append(promoted_line)
                        print(f"HEADING PROMOTED: '{promoted_line}' (was H{hash_count})")
                else:
                    heading_levels.append(None)  # Not a heading
                    lines.append(line)

            md_snippet = '\n'.join(lines[:10])
            print(f"MARKDOWN SNIPPET (first 10 lines):\n{md_snippet}")
            current_app.logger.info(f"Parsed Markdown snippet for {imp_doc.filename}:\n{md_snippet}")
        except Exception as e:
            # Fallback: attempt to read headings and paragraphs via python-docx
            print(f"PANDOC/CONVERSION ERROR: {e}")
            current_app.logger.error(f"Pandoc or conversion failed for {imp_doc.filename}: {e}")
            snippet = file_content[:500]
            current_app.logger.error(f"First 500 bytes of file: {snippet}")

            try:
                doc = Document(io.BytesIO(file_content))
                # Extract paragraphs and headings
                for p in doc.paragraphs:
                    text = (p.text or '').strip()
                    if not text:
                        continue
                    try:
                        ps = getattr(p, 'style', None)
                        style_name_raw = getattr(ps, 'name', '') or ''
                    except Exception:
                        style_name_raw = ''
                    style_name = style_name_raw.lower()

                    level = detect_heading_level_from_style(style_name)
                    if level is not None and style_name and not re.search(r'heading\s*'+str(level), style_name):
                        # Log only when mapping from a non-standard variant
                        print(f"STYLE->HEADING: '{style_name_raw}' -> H{level}")
                    if level == 1:
                        lines.append(f"# {text}")
                    elif level and level > 1:
                        hashes = '#' * min(level, 6)
                        lines.append(f"{hashes} {text}")
                    else:
                        lines.append(text)
                # Extract tables and convert to Markdown
                for table in doc.tables:
                    # Get all rows as lists of cell text
                    table_rows = []
                    for row in table.rows:
                        table_rows.append([cell.text.strip().replace('\n', ' ') for cell in row.cells])
                    if not table_rows:
                        continue
                    # Build Markdown pipe table
                    header = table_rows[0]
                    aligns = ['---'] * len(header)
                    md_table = ['| ' + ' | '.join(header) + ' |', '| ' + ' | '.join(aligns) + ' |']
                    for row in table_rows[1:]:
                        md_table.append('| ' + ' | '.join(row) + ' |')
                    lines.append('')
                    lines.extend(md_table)
                    lines.append('')
                md_snippet = '\n'.join(lines[:10])
                print(f"DOCX FALLBACK SNIPPET (first 10 lines):\n{md_snippet}")
                current_app.logger.info(f"DOCX fallback parsed snippet for {imp_doc.filename}:\n{md_snippet}")
            except Exception as e2:
                print(f"DOCX FALLBACK ERROR: {e2}")
                current_app.logger.error(f"DOCX fallback failed for {imp_doc.filename}: {e2}")
                lines = []

        paras = [('md', line) for line in lines]
        full_text = '\n'.join(lines)
    else:
        # Markdown file processing with image validation
        raw = file.read().decode('utf-8')
        
        # For markdown files, validate existing image references
        if imp_doc.id:  # Only if document is already saved
            image_handler = ImageHandler(imp_doc.id)
            validation_issues = image_handler.validate_markdown_images(raw)
            if validation_issues:
                for issue in validation_issues:
                    print(f"MARKDOWN IMAGE VALIDATION: {issue['message']}")
        
        # Optionally promote headers based on preserve_hierarchy setting
        lines = []
        for line in raw.splitlines():
            if line.strip().startswith('#'):
                # Count the number of # characters
                hash_count = len(line) - len(line.lstrip('#'))
                if not preserve_hierarchy and hash_count > 1:
                    # Promote to H1: replace multiple # with single #
                    content = line.lstrip('#').strip()
                    line = f"# {content}"
                    print(f"PROMOTED: '{line.strip()}' (was H{hash_count})")
                elif preserve_hierarchy:
                    print(f"PRESERVED: '{line.strip()}' (level H{hash_count})")
            lines.append(line)
        
        paras = [('md', line) for line in lines]
        full_text = '\n'.join(lines)

    items, buffer, order, current_title = [], [], 0, None
    print(f"PARSING: source={source}, paragraphs={len(paras)}")

    def commit_buffer():
        nonlocal order, current_title, buffer
        if current_title:
            content = '\n'.join(buffer).strip()
            
            # Apply additional content cleaning
            if source == 'word':
                # For Word documents, remove ALL blank lines as requested
                content = _remove_all_blank_lines(content)
            else:
                # For Markdown documents, use the existing cleaning that preserves paragraph breaks
                content = _clean_topic_content(content)
            
            # Only create an item if we have actual content (not just empty lines/whitespace)
            if content:
                items.append((order, current_title, content))
                print(f"COMMITTED: order={order}, title='{current_title}', content_len={len(content)}")
                order += 1
            else:
                print(f"SKIPPED EMPTY: title='{current_title}' (no substantive content)")
            buffer = []

    for style, text in paras:
        if preserve_hierarchy:
            # With hierarchy preservation, any heading level can be a topic
            is_heading = text.strip().startswith('#')
        else:
            # Original behavior: only H1 creates new topics
            is_heading = text.strip().startswith('#') and not text.strip().startswith('##')
        
        print(f"LINE: '{text}' -> heading={is_heading} (preserve_hierarchy={preserve_hierarchy})")
        
        if is_heading:
            # Check if we have a current title but no substantive content yet
            current_buffer_content = '\n'.join(buffer).strip()
            current_buffer_has_content = bool(current_buffer_content and 
                                            not all(line.strip() == '' or line.strip().startswith('#') 
                                                   for line in current_buffer_content.split('\n')))
            
            if current_title and not current_buffer_has_content:
                # Merge this heading into the content of the previous heading
                heading_text = text.strip().lstrip('#').strip()
                if preserve_hierarchy:
                    # Keep original heading level when merging
                    hash_count = len(text.strip()) - len(text.strip().lstrip('#'))
                    hashes = '#' * min(hash_count + 1, 6)  # Make it one level deeper in content
                    buffer.append(f"{hashes} {heading_text}")
                else:
                    buffer.append(f"## {heading_text}")  # Add as H2 in content
                print(f"MERGED_HEADING: '{heading_text}' added to content of '{current_title}' (no substantive content found)")
            else:
                # Normal case: commit previous section and start new one
                commit_buffer()
                current_title = text.strip().lstrip('#').strip()
                print(f"NEW_TITLE: '{current_title}'")
        else:
            buffer.append(text)

    commit_buffer()
    print(f"FINAL: {len(items)} items created")

    # Fallback: if no items were created, use entire content as a single item
    if not items:
        try:
            fallback_title = os.path.splitext(imp_doc.filename)[0] if getattr(imp_doc, 'filename', None) else 'Imported Document'
            fallback_content = full_text if 'full_text' in locals() else ''
            if source == 'word':
                fallback_content = _remove_all_blank_lines(fallback_content)
            else:
                fallback_content = _clean_topic_content(fallback_content)

            if fallback_content and fallback_content.strip():
                print(f"FALLBACK: Creating single item from full document. title='{fallback_title}', content_len={len(fallback_content)}")
                items.append((0, fallback_title, fallback_content))
            else:
                print("FALLBACK: Full document content is empty after cleaning; no item created")
        except Exception as e:
            print(f"FALLBACK ERROR: {e}")

    # Extract links from the full document content
    total_links_extracted = 0
    try:
        links_count = _extract_and_store_links(imp_doc.id, full_text)
        total_links_extracted += links_count
        print(f"LINK EXTRACTION: Extracted {links_count} links from document")
    except Exception as e:
        print(f"LINK EXTRACTION ERROR: {e}")
        current_app.logger.error(f"Link extraction failed for {imp_doc.filename}: {e}")

    for order, title, content in items:
        # If preserve_hierarchy is enabled and we have heading level info, encode it in the title
        if preserve_hierarchy and len(items) > order:
            # Find the original heading level for this title by checking the parsed lines
            heading_level = 1  # Default to H1
            for line in full_text.split('\n'):
                if line.strip().lstrip('#').strip() == title.strip():
                    heading_level = len(line) - len(line.lstrip('#')) if line.strip().startswith('#') else 1
                    break
            # Encode level in title format: "LEVEL:3:Actual Title"
            encoded_title = f"LEVEL:{heading_level}:{title}"
        else:
            encoded_title = title
            
        db.session.add(ImportItem(
            document_id=imp_doc.id,
            heading_order=order,
            title=encoded_title,
            content=content
        ))
    
    # Store the link extraction summary in logs
    if total_links_extracted > 0:
        current_app.logger.info(f"Successfully extracted {total_links_extracted} links from {imp_doc.filename}")
    
    return total_links_extracted


def _upload_file(source):
    print(f"UPLOAD: Starting upload with source={source}")
    file = request.files.get('file')
    import_type = request.form.get('import_type', 'topics')  # Default to topics for backward compatibility
    preserve_hierarchy = request.form.get('preserve_hierarchy', 'false').lower() == 'true'
    
    if not file or source not in SOURCES:
        print(f"UPLOAD: Missing file or invalid source. file={file}, source={source}")
        return jsonify({'error': 'Missing file or invalid source'}), 400

    try:
        # Handle collection import
        if import_type == 'collection':
            return _import_as_collection(file, source)
        else:
            # Handle regular topic import
            return _import_as_topics(file, source, preserve_hierarchy)

    except Exception as e:
        db.session.rollback()
        current_app.logger.exception("Upload failed")
        print(f"UPLOAD: Exception occurred: {e}")
        return jsonify({'error': str(e)}), 500


def _import_as_topics(file, source, preserve_hierarchy=False):
    """Import document as individual topics (original functionality)
    
    Args:
        file: The uploaded file
        source: The source type (e.g., 'word', 'text')
        preserve_hierarchy: If True, preserve heading hierarchy by creating a collection; if False, promote all to H1
    """
    if preserve_hierarchy:
        # When hierarchy preservation is requested, automatically create as collection
        # Use package-relative import to avoid importing backend.models twice
        from ..models import Collection, Topic, collection_topic_tree, Project
        
        # Set default collection parameters
        import datetime
        collection_name = f"Document Import - {secure_filename(file.filename)}"
        collection_form_number = f"AUTO_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Try to get a default project (use the first available project)
        default_project = Project.query.first()
        if not default_project:
            return jsonify({'error': 'No projects available. Please create a project first before importing with hierarchy.'}), 400
        
        # Create a temporary ImportDocument for proper image processing
        temp_imp_doc = ImportDocument(filename=secure_filename(file.filename), source_type=source)
        db.session.add(temp_imp_doc)
        db.session.flush()  # get temp_imp_doc.id for image processing
        
        # Parse document with hierarchical structure preservation AND image processing
        hierarchical_items = _parse_hierarchical_structure_with_images(file, source, temp_imp_doc.id)
        
        if not hierarchical_items:
            print("HIERARCHICAL PARSING FAILED: Falling back to regular flat parsing")
            # Fall back to regular parsing logic
            file.stream.seek(0)
            temp_imp_doc = ImportDocument(filename=secure_filename(file.filename), source_type=source)
            db.session.add(temp_imp_doc)
            db.session.flush()
            
            _parse_and_store(file, temp_imp_doc, source, preserve_hierarchy=False)
            items_count = ImportItem.query.filter_by(document_id=temp_imp_doc.id).count()
            
            if items_count == 0:
                # Commit any extracted links before rolling back
                links_count = ImportLink.query.filter_by(document_id=temp_imp_doc.id).count()
                print(f"COLLECTION IMPORT: Preserving {links_count} extracted links before rollback")
                if links_count > 0:
                    db.session.commit()
                
                db.session.rollback()
                error_msg = f"No content items could be extracted from the document. "
                if source == 'word':
                    error_msg += "This may be due to: 1) The document has no recognizable headings, 2) Pandoc conversion failed, or 3) The document structure is not supported."
                else:
                    error_msg += "This may be due to: 1) The document has no H1 headings (# Title), or 2) The file is empty or corrupted."
                return jsonify({'error': error_msg}), 422
            
            # Convert ImportItems to hierarchical format
            import_items = ImportItem.query.filter_by(document_id=temp_imp_doc.id).order_by(ImportItem.heading_order).all()
            hierarchical_items = []
            for item in import_items:
                hierarchical_items.append({
                    'title': item.title,
                    'content': item.content,
                    'level': 1,  # All items are H1 in flat parsing
                    'parent_index': None  # No hierarchy in fallback
                })
            
            # Clean up temporary import document
            ImportItem.query.filter_by(document_id=temp_imp_doc.id).delete()
            db.session.delete(temp_imp_doc)
            
            print(f"FALLBACK SUCCESS: Created {len(hierarchical_items)} flat items")
        
        # Create the collection
        collection = Collection(
            name=collection_name,
            form_number=collection_form_number,
            description=f"Auto-generated collection from {file.filename} with preserved hierarchy",
            project_id=default_project.id
        )
        collection.archived = False
        db.session.add(collection)
        db.session.flush()  # get collection.id
        
        # Create topics with hierarchy - do this in two passes
        created_topics = []
        topic_id_map = {}
        
        # Pass 1: Create all topics first
        for i, item in enumerate(hierarchical_items):
            # Clean the content
            content = _remove_all_blank_lines(item['content']) if item['content'] else ''
            
            topic = Topic(
                title=item['title'],
                content=content
            )
            db.session.add(topic)
            db.session.flush()  # get topic.id
            created_topics.append(topic)
            topic_id_map[i] = topic.id
        
        # Pass 2: Create hierarchical relationships
        for i, item in enumerate(hierarchical_items):
            # Determine parent topic ID
            parent_topic_id = None
            if item['parent_index'] is not None and item['parent_index'] in topic_id_map:
                parent_topic_id = topic_id_map[item['parent_index']]
            
            # Add topic to collection with hierarchical relationship
            db.session.execute(
                collection_topic_tree.insert().values(
                    collection_id=collection.id,
                    topic_id=topic_id_map[i],
                    parent_topic_id=parent_topic_id,
                    position=i
                )
            )
        
        db.session.commit()
        print(f"HIERARCHICAL_TOPIC_IMPORT: Created collection '{collection_name}' with {len(created_topics)} topics")
        
        # Return collection info instead of import document
        collection_dict = collection.to_dict()
        collection_dict['topics_count'] = len(created_topics)
        collection_dict['message'] = f"Document imported as collection with {len(created_topics)} topics in hierarchical structure"
        
        return jsonify(collection_dict), 201
    
    # Original flat topic import logic
    imp_doc = ImportDocument(
        filename=secure_filename(file.filename),
        source_type=source
    )
    db.session.add(imp_doc)
    db.session.flush()  # get imp_doc.id
    print(f"UPLOAD: Created ImportDocument with ID={imp_doc.id}")

    _parse_and_store(file, imp_doc, source, preserve_hierarchy)
    
    # Check if any items were created
    items_count = ImportItem.query.filter_by(document_id=imp_doc.id).count()
    print(f"UPLOAD: Created {items_count} import items")
    
    if items_count == 0:
        # Commit any extracted links so they aren't lost when we delete the import document
        links_count = ImportLink.query.filter_by(document_id=imp_doc.id).count()
        print(f"UPLOAD: Preserving {links_count} extracted links before deleting import")
        if links_count > 0:
            db.session.commit()
        
        # Delete the ImportDocument since it has no items, but keep the links
        db.session.delete(imp_doc)
        db.session.commit()
        # Try to log more details for debugging
        file_content = b""  # ensure defined even if read fails below
        try:
            file.stream.seek(0)
            file_content = file.read()
            snippet = file_content[:500]
        except Exception as e:
            snippet = f"Could not read file content: {e}"
        # Try to get a snippet of the Markdown if possible
        try:
            markdown_content = _convert_word_to_markdown(file_content, imp_doc.id) if source == 'word' else file_content.decode('utf-8')
            md_lines = markdown_content.splitlines()[:10]
            md_snippet = '\n'.join(md_lines)
        except Exception as e:
            md_snippet = f"Could not get Markdown snippet: {e}"
        error_msg = f"No content items could be extracted from {imp_doc.filename}. "
        if source == 'word':
            error_msg += "This may be due to: 1) The document has no recognizable headings, 2) Pandoc conversion failed, or 3) The document structure is not supported."
        else:
            error_msg += "This may be due to: 1) The document has no H1 headings (# Title), or 2) The file is empty or corrupted."
        error_msg += f"\nFirst 500 bytes of file: {snippet}\nMarkdown snippet (first 10 lines):\n{md_snippet}"
        print(f"UPLOAD: {error_msg}")
        current_app.logger.error(error_msg)
        return jsonify({'error': error_msg}), 422
    
    db.session.commit()
    print(f"UPLOAD: Committed to database")
    return jsonify(imp_doc.to_dict(include_items=True)), 201


def _parse_hierarchical_structure_with_images(file, source, import_doc_id):
    """Parse document preserving hierarchical heading structure WITH image processing"""
    file.stream.seek(0)
    
    try:
        # Get markdown content based on source type WITH proper image processing
        if source == 'word':
            file_content = file.read()
            file.stream.seek(0)
            # Use full image processing for hierarchical parsing
            markdown_content = _convert_word_to_markdown(file_content, import_doc_id)
            # Extract links from converted markdown and store under ImportDocument
            try:
                links_count = _extract_and_store_links(import_doc_id, markdown_content)
                current_app.logger.info(f"HIERARCHICAL IMPORT: Extracted {links_count} links for import document {import_doc_id}")
            except Exception as e:
                current_app.logger.error(f"HIERARCHICAL IMPORT: Link extraction failed for import document {import_doc_id}: {e}")
        else:
            # For markdown files, read directly
            markdown_content = file.read().decode('utf-8')
            file.stream.seek(0)
            
            # For markdown files, validate existing image references
            if import_doc_id:
                image_handler = ImageHandler(import_doc_id)
                validation_issues = image_handler.validate_markdown_images(markdown_content)
                if validation_issues:
                    for issue in validation_issues:
                        print(f"MARKDOWN IMAGE VALIDATION: {issue['message']}")
            # Extract links from markdown input and store
            try:
                links_count = _extract_and_store_links(import_doc_id, markdown_content)
                current_app.logger.info(f"HIERARCHICAL IMPORT (markdown): Extracted {links_count} links for import document {import_doc_id}")
            except Exception as e:
                current_app.logger.error(f"HIERARCHICAL IMPORT (markdown): Link extraction failed for import document {import_doc_id}: {e}")

        return _parse_hierarchical_content(markdown_content)
        
    except Exception as e:
        print(f"HIERARCHICAL PARSING ERROR: {e}")
        return []


def _parse_hierarchical_content(markdown_content):
    """Parse markdown content into hierarchical structure"""
    hierarchical_items = []
    current_stack = []  # Stack to track heading hierarchy levels
    current_content = []
    
    for line in markdown_content.splitlines():
        stripped = line.strip()
        
        if stripped.startswith('#'):
            # This is a heading - determine its level
            hash_count = len(line) - len(line.lstrip('#'))
            title = stripped.lstrip('#').strip()
            
            # Commit content to the current item in stack
            if current_stack and current_content:
                content_text = '\n'.join(current_content).strip()
                if content_text:
                    current_stack[-1]['content'] = content_text
                current_content = []
            
            # Pop items from stack that are at same or deeper level
            while current_stack and current_stack[-1]['level'] >= hash_count:
                completed_item = current_stack.pop()
                hierarchical_items.append(completed_item)
            
            # Create new heading item
            heading_item = {
                'title': title,
                'level': hash_count,
                'content': '',
                'parent_index': None
            }
            
            # Set parent reference if there's a parent in the stack
            if current_stack:
                # The parent is the most recent item at a shallower level (stack top)
                parent_item = current_stack[-1]
                # We'll set the parent_index after we know where the parent will be in the final list
                heading_item['parent_item'] = parent_item
            
            current_stack.append(heading_item)
        else:
            # This is content - add to current content buffer
            current_content.append(line)
    
    # Don't forget content for the last heading(s)
    if current_stack and current_content:
        content_text = '\n'.join(current_content).strip()
        if content_text:
            current_stack[-1]['content'] = content_text
    
    # Pop all remaining items from stack
    while current_stack:
        completed_item = current_stack.pop()
        hierarchical_items.append(completed_item)
    
    # Now resolve parent indices correctly
    # Since items are added to the list when they're popped from the stack,
    # parents appear AFTER children in the list. We need to reverse the lookup.
    for i, item in enumerate(hierarchical_items):
        if 'parent_item' in item:
            parent_item = item['parent_item']
            # Find the parent in the hierarchical_items list (it should appear after this item)
            for j in range(i + 1, len(hierarchical_items)):
                potential_parent = hierarchical_items[j]
                if (potential_parent['title'] == parent_item['title'] and 
                    potential_parent['level'] == parent_item['level']):
                    item['parent_index'] = j
                    break
            # Remove the temporary parent_item reference
            del item['parent_item']
        else:
            item['parent_index'] = None
    
    print(f"HIERARCHICAL PARSING: Found {len(hierarchical_items)} hierarchical items")
    return hierarchical_items


def _parse_hierarchical_structure(file, source):
    """Parse document preserving hierarchical heading structure (legacy version without images)"""
    file.stream.seek(0)
    
    try:
        # Get markdown content based on source type
        if source == 'word':
            file_content = file.read()
            file.stream.seek(0)
            # For hierarchical parsing, skip image processing to avoid DB constraint issues
            # Images will be processed later when the real document is created
            markdown_content = _convert_word_to_markdown_no_images(file_content)
        else:
            # For markdown files, read directly
            markdown_content = file.read().decode('utf-8')
            file.stream.seek(0)
        
        return _parse_hierarchical_content(markdown_content)
        
        hierarchical_items = []
        current_stack = []  # Stack to track heading hierarchy levels
        current_content = []
        
        for line in markdown_content.splitlines():
            stripped = line.strip()
            
            if stripped.startswith('#'):
                # This is a heading - determine its level
                hash_count = len(line) - len(line.lstrip('#'))
                title = stripped.lstrip('#').strip()
                
                # Commit content to the current item in stack
                if current_stack and current_content:
                    content_text = '\n'.join(current_content).strip()
                    if content_text:
                        current_stack[-1]['content'] = content_text
                    current_content = []
                
                # Pop items from stack that are at same or deeper level
                while current_stack and current_stack[-1]['level'] >= hash_count:
                    completed_item = current_stack.pop()
                    hierarchical_items.append(completed_item)
                
                # Create new heading item
                heading_item = {
                    'title': title,
                    'level': hash_count,
                    'content': '',
                    'parent_index': None
                }
                
                # Set parent relationship - we'll fix this after all items are processed
                # For now, just track the parent level
                heading_item['parent_level'] = current_stack[-1]['level'] if current_stack else None
                
                current_stack.append(heading_item)
                print(f"HIERARCHICAL HEADING: Level {hash_count} - '{title}' (parent: {heading_item['parent_index']})")
                
            else:
                # Regular content line
                current_content.append(line)
        
        # Process remaining items in stack
        if current_stack and current_content:
            content_text = '\n'.join(current_content).strip()
            if content_text:
                current_stack[-1]['content'] = content_text
        
        # Add all remaining stack items to hierarchical_items
        hierarchical_items.extend(current_stack)
        
        # Fix parent relationships - find the correct parent for each item
        for i, item in enumerate(hierarchical_items):
            if item.get('parent_level') is not None:
                # Find the most recent item with the parent level
                for j in range(i - 1, -1, -1):
                    if hierarchical_items[j]['level'] == item['parent_level']:
                        item['parent_index'] = j
                        break
            # Remove the temporary parent_level field
            item.pop('parent_level', None)
        
        print(f"HIERARCHICAL PARSING: Created {len(hierarchical_items)} hierarchical items")
        return hierarchical_items
        
    except Exception as e:
        print(f"HIERARCHICAL PARSING ERROR: {e}")
        current_app.logger.error(f"Hierarchical parsing failed: {e}")
        # Return empty list to trigger the error handling in calling function
        return []


def _import_as_collection(file, source):
    """Import document as a collection with hierarchical structure"""
    # Use package-relative import to avoid importing backend.models twice
    from ..models import Collection, Topic, collection_topic_tree, Project, ImportDocument
    from werkzeug.utils import secure_filename
    
    # Get collection details from form
    collection_name = request.form.get('collection_name', '').strip()
    collection_form_number = request.form.get('collection_form_number', '').strip()
    collection_description = request.form.get('collection_description', '').strip()
    project_id = request.form.get('project_id', '').strip()
    
    if not collection_name:
        return jsonify({'error': 'Collection name is required'}), 400
    if not collection_form_number:
        return jsonify({'error': 'Collection ID (Form Number) is required'}), 400
    if not project_id:
        return jsonify({'error': 'Project selection is required'}), 400
    
    # Validate that the project exists
    try:
        project_id = int(project_id)
        project = Project.query.get(project_id)
        if not project:
            return jsonify({'error': 'Selected project does not exist'}), 400
    except ValueError:
        return jsonify({'error': 'Invalid project ID'}), 400
    
    # Check if form number already exists
    existing_collection = Collection.query.filter_by(form_number=collection_form_number).first()
    if existing_collection:
        return jsonify({'error': f'Collection ID "{collection_form_number}" already exists'}), 400
    
    # Create a temporary ImportDocument for proper image processing
    temp_imp_doc = ImportDocument(filename=secure_filename(file.filename), source_type=source)
    db.session.add(temp_imp_doc)
    db.session.flush()  # get temp_imp_doc.id for image processing

    # Parse document with hierarchical structure preservation AND image processing
    hierarchical_items = _parse_hierarchical_structure_with_images(file, source, temp_imp_doc.id)

    if not hierarchical_items:
        error_msg = f"No content items could be extracted from the document. "
        if source == 'word':
            error_msg += "This may be due to: 1) The document has no recognizable headings, 2) Pandoc conversion failed, or 3) The document structure is not supported."
        else:
            error_msg += "This may be due to: 1) The document has no H1 headings (# Title), or 2) The file is empty or corrupted."
        return jsonify({'error': error_msg}), 422
    
    print(f"COLLECTION_IMPORT: Parsed {len(hierarchical_items)} hierarchical items from document")
    
    # Create the collection
    collection = Collection(
        name=collection_name,
        form_number=collection_form_number,
        description=collection_description or None,
        project_id=project_id
    )
    db.session.add(collection)
    db.session.flush()  # get collection.id
    print(f"COLLECTION_IMPORT: Created collection with ID={collection.id}")
    
    # Convert hierarchical items to topics and add them to the collection with proper hierarchy
    # IMPORTANT: Parents appear AFTER children in hierarchical_items due to stack-pop order during parsing.
    # Therefore we must do this in TWO PASSES: (1) create topics, (2) create relationships using parent_index.
    created_topics = []
    topic_id_map = {}  # Map from item index to topic ID for parent relationships

    # Pass 1: Create all topics first so every index is available
    for i, item in enumerate(hierarchical_items):
        content = _remove_all_blank_lines(item['content']) if item['content'] else ''
        topic = Topic(title=item['title'], content=content)
        db.session.add(topic)
        db.session.flush()  # get topic.id
        created_topics.append(topic)
        topic_id_map[i] = topic.id
        print(f"COLLECTION_IMPORT: Created topic '{item['title']}' (level {item['level']}) -> topic_id={topic.id}")

    # Pass 2: Insert hierarchical relationships using parent_index mapping
    linked = 0
    for i, item in enumerate(hierarchical_items):
        parent_topic_id = None
        if item['parent_index'] is not None:
            parent_idx = item['parent_index']
            parent_topic_id = topic_id_map.get(parent_idx)

        db.session.execute(
            collection_topic_tree.insert().values(
                collection_id=collection.id,
                topic_id=topic_id_map[i],
                parent_topic_id=parent_topic_id,
                position=i
            )
        )

        hierarchy_info = f"parent: {parent_topic_id}" if parent_topic_id else "root level"
        print(f"COLLECTION_IMPORT: Linked topic_idx={i} -> {hierarchy_info}")
        if parent_topic_id:
            linked += 1

    print(f"COLLECTION_IMPORT: Created {len(created_topics)} topics; linked {linked} with parents")
    
    # Commit everything
    db.session.commit()
    print(f"COLLECTION_IMPORT: Successfully committed collection import and cleaned up temporary data")
    
    return jsonify({
        'collection_id': collection.id,
        'collection_name': collection.name,
        'collection_form_number': collection.form_number,
        'topics_count': len(created_topics),
        'message': f'Successfully imported {len(created_topics)} topics into collection "{collection.name}"'
    }), 201


@import_bp.route('/upload', methods=['POST'])
def upload_generic():
    return _upload_file(request.form.get('source', '').lower())


@import_bp.route('/markdown', methods=['POST'])
def upload_markdown():
    return _upload_file('markdown')


@import_bp.route('/history', methods=['GET'])
def get_import_history():
    """Get list of all import documents with their status"""
    try:
        current_app.logger.info("Fetching import history...")
        docs = ImportDocument.query.order_by(ImportDocument.created_at.desc()).all()
        current_app.logger.info(f"Found {len(docs)} import documents")
        
        result = []
        for i, doc in enumerate(docs):
            try:
                current_app.logger.info(f"Processing doc {i+1}: ID={doc.id}, filename={doc.filename}")
                doc_dict = doc.to_dict(include_items=False)
                result.append(doc_dict)
                current_app.logger.info(f"Successfully processed doc {doc.id}")
            except Exception as e:
                current_app.logger.error(f"Error processing doc {doc.id}: {str(e)}")
                # Continue with other documents instead of failing completely
                continue
                
        current_app.logger.info(f"Returning {len(result)} documents")
        return jsonify(result), 200
    except Exception as e:
        current_app.logger.error(f"Error fetching import history: {str(e)}")
        import traceback
        current_app.logger.error(traceback.format_exc())
        return jsonify({'error': f'Failed to fetch import history: {str(e)}'}), 500


@import_bp.route('/staging/<int:doc_id>/images', methods=['GET'])
def get_import_images(doc_id):
    """Get all images associated with an import document"""
    try:
        doc = ImportDocument.query.get_or_404(doc_id)
        
        # Get images from database
        db_images = ImportImage.query.filter_by(document_id=doc_id).all()
        
        # Validate that files actually exist on disk
        # Only return images whose files exist to prevent 404s on the frontend
        from pathlib import Path
        validated_images = []
        for img in db_images:
            backend_path = Path(img.backend_path)
            frontend_path = Path(img.frontend_path)
            # Check if file exists in either backend or frontend location
            if backend_path.exists() or frontend_path.exists():
                validated_images.append(img.to_dict())
            else:
                current_app.logger.warning(
                    f"Image file missing for {doc_id}/{img.filename}: "
                    f"backend={backend_path.exists()}, frontend={frontend_path.exists()}"
                )
        
        # Also get images from filesystem (in case of sync issues)
        image_handler = ImageHandler(doc_id)
        fs_images = image_handler.get_import_images()
        
        # If database is out of sync with filesystem, use filesystem as source of truth
        if len(validated_images) == 0 and len(fs_images) > 0:
            current_app.logger.info(
                f"No valid database images for import {doc_id} but found {len(fs_images)} on filesystem; "
                f"using filesystem as source of truth"
            )
            images_data = fs_images
        else:
            images_data = validated_images
        
        return jsonify({
            'document_id': doc_id,
            'document_filename': doc.filename,
            'images': images_data,
            'filesystem_images': fs_images,
            'total_count': len(images_data)
        }), 200
        
    except Exception as e:
        current_app.logger.error(f"Error fetching images for import {doc_id}: {str(e)}")
        return jsonify({'error': f'Failed to fetch images: {str(e)}'}), 500


@import_bp.route('/staging/<int:doc_id>/links', methods=['GET'])
def get_import_links(doc_id):
    """Get all links extracted from an import document"""
    try:
        doc = ImportDocument.query.get_or_404(doc_id)
        
        # Get links from database
        db_links = ImportLink.query.filter_by(document_id=doc_id).order_by(ImportLink.position_in_document).all()
        links_data = [link.to_dict() for link in db_links]
        
        return jsonify({
            'document_id': doc_id,
            'document_filename': doc.filename,
            'links': links_data,
            'total_count': len(links_data)
        }), 200
        
    except Exception as e:
        current_app.logger.error(f"Error fetching links for import {doc_id}: {str(e)}")
        return jsonify({'error': f'Failed to fetch links: {str(e)}'}), 500


@import_bp.route('/staging/<int:doc_id>', methods=['GET'])
def get_staging(doc_id):
    doc = ImportDocument.query.get_or_404(doc_id)
    result = doc.to_dict(include_items=True)
    
    # Include image information
    images = ImportImage.query.filter_by(document_id=doc_id).all()
    result['images'] = [img.to_dict() for img in images]
    result['images_count'] = len(images)
    
    return jsonify(result), 200


@import_bp.route('/staging/<int:doc_id>/sme_approve', methods=['POST'])
def sme_approve(doc_id):
    """Approve an import document by SME"""
    try:
        doc = ImportDocument.query.get_or_404(doc_id)
        doc.status = 'approved'  # Changed from 'sme_approved' to 'approved'
        doc.review_step = 'sme_approved'  # Set the review step for frontend
        db.session.commit()
        current_app.logger.info(f"Import document {doc_id} approved by SME")
        return jsonify({'message': 'Import approved successfully'}), 200
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error approving import {doc_id}: {str(e)}")
        return jsonify({'error': 'Failed to approve import'}), 500


@import_bp.route('/staging/<int:doc_id>/commit', methods=['POST'])
def commit_import(doc_id):
    """Commit an approved import to create actual topics"""
    try:
        doc = ImportDocument.query.get_or_404(doc_id)
        if doc.status != 'approved':  # Changed from 'sme_approved' to 'approved'
            return jsonify({'error': 'Import must be approved before commit'}), 400
        
        # Create topics from import items
        for item in doc.items:
            # Check if this item has encoded level information and decode it
            title = item.title
            if title.startswith("LEVEL:"):
                # Decode level information: "LEVEL:3:Actual Title"
                parts = title.split(":", 2)
                if len(parts) == 3:
                    title = parts[2]  # Extract the actual title
            
            # Create topic from import item (Topic has no heading_order field)
            topic = Topic(
                title=title,
                content=item.content
            )
            db.session.add(topic)
        
        # Update status and review step to show final approval
        doc.review_step = 'final_approved'
        db.session.commit()
        current_app.logger.info(f"Import document {doc_id} committed successfully")
        return jsonify({'message': 'Import committed successfully'}), 200
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error committing import {doc_id}: {str(e)}")
        return jsonify({'error': 'Failed to commit import'}), 500


@import_bp.route('/staging/<int:doc_id>/reprocess', methods=['POST'])
def reprocess_document(doc_id):
    """Reprocess an existing import document to extract items"""
    try:
        doc = ImportDocument.query.get_or_404(doc_id)
        
        # Delete existing items
        ImportItem.query.filter_by(document_id=doc_id).delete()
        
        # Since we don't have the original file, we can't reprocess
        # This would need to be enhanced to store the original file
        return jsonify({'error': 'Reprocessing requires the original file to be stored'}), 400
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error reprocessing document {doc_id}: {str(e)}")
        return jsonify({'error': f'Failed to reprocess document: {str(e)}'}), 500


@import_bp.route('/staging/<int:doc_id>/reject', methods=['POST'])
def reject_import(doc_id):
    """Reject an import document"""
    try:
        doc = ImportDocument.query.get_or_404(doc_id)
        doc.status = 'rejected'
        db.session.commit()
        current_app.logger.info(f"Import document {doc_id} rejected")
        return jsonify({'message': 'Import rejected successfully'}), 200
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error rejecting import {doc_id}: {str(e)}")
        return jsonify({'error': 'Failed to reject import'}), 500
